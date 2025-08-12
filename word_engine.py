from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_bg_color(cell, fill):
    """Set background color of a cell (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def generate_memo():
    doc = Document()

    # Variables
    variables = {
        "name": "Roman",
        "city": "New York",
        "role": "Sales Lead"
    }

    # Intro paragraph
    text = "Hello, my name is {name}. I live in {city}, and I work as a {role}."
    doc.add_paragraph(text.format(**variables))

    # ---- CREATE A NICELY-STYLED TABLE ----
    data = [
        ["Month", "Revenue", "Expenses", "Profit"],
        ["January", "$50,000", "$30,000", "$20,000"],
        ["February", "$55,000", "$32,000", "$23,000"],
        ["March", "$60,000", "$35,000", "$25,000"]
    ]

    table = doc.add_table(rows=0, cols=len(data[0]))
    table.style = "Light List Accent 1"  # clean built-in style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- HEADER ROW ---
    hdr_cells = table.add_row().cells
    for idx, heading in enumerate(data[0]):
        hdr_cells[idx].text = heading
        for p in hdr_cells[idx].paragraphs:
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_cell_bg_color(hdr_cells[idx], "2F5496")  # dark blue

    # --- DATA ROWS ---
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            row_cells[idx].text = value
            for p in row_cells[idx].paragraphs:
                run = p.runs[0]
                run.font.size = Pt(11)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.alignment = WD_TABLE_ALIGNMENT.CENTER if idx > 0 else WD_TABLE_ALIGNMENT.LEFT
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.save("generated_memo.docx")

generate_memo()
