from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

def generate_word_from_html(html_content: str, output_path: str = "memo.docx"):
    """
    Converts HTML content into a Word document.
    Handles paragraphs and tables inside divs.
    """
    doc = Document()
    soup = BeautifulSoup(html_content, "html.parser")

    # Process all divs recursively
    for div in soup.find_all("div", recursive=False):
        for child in div.children:
            if child.name == "div":
                # Nested divs (section titles or paragraphs)
                text = child.get_text(strip=True)
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(6)
            elif child.name == "table":
                rows = child.find_all("tr")
                if rows:
                    word_table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["td","th"])))
                    word_table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        cells = row.find_all(["td","th"])
                        for j, cell in enumerate(cells):
                            word_table.cell(i,j).text = cell.get_text()
            elif child.name is None:
                # Direct text node
                text = str(child).strip()
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    return output_path
